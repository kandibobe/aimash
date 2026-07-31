"""Сборка `ТЗ.md` из трёх оригиналов заказчика (.docx) — источник истины в читаемом виде.

Зачем скрипт, а не разовая конвертация. `CLAUDE.md` объявляет `ТЗ.md` источником истины, но
файла не было ни в дереве, ни в истории git: сверять код с ТЗ было физически нечем, а
единственным мостом до оригиналов служила производная матрица §33 `deploy/hermes/HERMES_SPEC.md`.
Скрипт делает мост воспроизводимым: заказчик прислал новую редакцию .docx — регенерируем.

Дрейф ловит `tests/test_tz_sync.py`: sha256 каждого источника записан в шапку `ТЗ.md`, и если
.docx обновили без регенерации, тест краснеет. Поэтому ручная вычитка markdown-разметки
безопасна — она хэши источников не трогает.

Гоча python-docx: `Document.paragraphs` и `Document.tables` — два независимых списка, порядок
«абзац → таблица → абзац» из них не восстанавливается. Идём по `document.element.body` и
различаем `w:p` / `w:tbl` вручную.

    python scripts/docx_to_tz.py           # перегенерировать ТЗ.md
    python scripts/docx_to_tz.py --check   # только проверить актуальность (для CI/теста)
"""

from __future__ import annotations

import hashlib
import sys
from pathlib import Path

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "ТЗ.md"

# Порядок = нумерация разделов, которой пользуется матрица трассируемости §33 HERMES_SPEC.
SOURCES: list[tuple[str, str, str]] = [
    ("Aimash_Technical_Specification.docx", "ТЗ-1", "Основное ТЗ, разделы 1–18"),
    (
        "Aimash_Flow_Google_Search_4.docx",
        "ТЗ-2",
        "Дополнение: раздел 19 — флоу создания Search-кампании",
    ),
    (
        "Информация о клиентах_1.docx",
        "ТЗ-3",
        "Дополнение: раздел 20 — база знаний клиента и краулинг",
    ),
]

HEADER = """# Aimash — Техническое задание (сводный текст оригиналов)

> ⚠️ **Файл сгенерирован** из трёх `.docx` заказчика скриптом `scripts/docx_to_tz.py`.
> Источник истины для этого архивного contract evidence — сами `.docx`; файл нужен, чтобы исходный
> заказ можно было читать и цитировать построчно. Правки смысла вносить сюда **нельзя** — только в `.docx`
> с последующей регенерацией. Допустима вычитка markdown-разметки (таблицы, заголовки,
> уровни списков), искажённой при конвертации из Word.
>
> Расхождение `.docx` с этим файлом ловит `tests/test_tz_sync.py` по sha256 ниже.
>
> 🕰️ **CONTRACT EVIDENCE, НЕ СПЕКА СБОРКИ.** Это дословный текст исходного заказа. Единственный
> текущий продуктовый канон — `SPEC.md`. Этот файл нужен только для сверки «что было заказано ↔
> что принято единым ТЗ» и не задаёт архитектуру или UX реализации.

## Источники

| Файл | Обозначение | Содержание | sha256 |
|---|---|---|---|
{sources_table}

Матрица соответствия исходных документов принятому продукту находится в `SPEC.md` §19.

---
"""


def sha256_of(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _list_level(p: Paragraph) -> int | None:
    """Уровень вложенности маркера списка или None. Тип маркера (цифра/точка) Word хранит в
    numbering.xml по numId — восстанавливать его не пытаемся, все списки выводим как `-`."""
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None:
        return None
    ilvl = pPr.numPr.ilvl
    return int(ilvl.val) if ilvl is not None and ilvl.val is not None else 0


def _bold_whole(p: Paragraph) -> bool:
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def para_md(p: Paragraph) -> str:
    text = p.text.strip()
    if not text:
        return ""
    style = p.style.name if p.style else ""
    if style.startswith("Heading"):
        tail = style.split()[-1]
        lvl = int(tail) if tail.isdigit() else 2
        return "#" * min(lvl + 1, 6) + " " + text
    lvl = _list_level(p)
    if lvl is not None:
        return "  " * lvl + "- " + text
    # Стиль не задан (в исходниках таких абзацев 161): короткая жирная строка без завершающей
    # пунктуации — почти всегда подзаголовок, остальное — обычный абзац.
    if _bold_whole(p) and len(text) < 90 and not text.endswith((".", ":", ";", "!", "?")):
        return "#### " + text
    return text


def table_md(t: Table) -> str:
    rows = [
        [c.text.strip().replace("\n", " ").replace("|", "\\|") for c in r.cells] for r in t.rows
    ]
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ""
    ncol = max(len(r) for r in rows)
    # Одноколоночная таблица в Word — это врезка/callout, а не данные: markdown-таблица из
    # одной колонки читается хуже цитаты.
    if ncol == 1:
        return "\n".join("> " + r[0] for r in rows if r and r[0])
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    out = ["| " + " | ".join(rows[0]) + " |", "|" + "---|" * ncol]
    out += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join(out)


def convert(path: Path) -> str:
    d = docx.Document(str(path))
    out: list[str] = []
    for child in d.element.body.iterchildren():
        if child.tag == W + "p":
            md = para_md(Paragraph(child, d))
            if md:
                out.append(md)
        elif child.tag == W + "tbl":
            md = table_md(Table(child, d))
            if md:
                out += ["", md, ""]
    lines: list[str] = []
    for line in out:  # схлопнуть подряд идущие пустые
        if line == "" and (not lines or lines[-1] == ""):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def build() -> str:
    missing = [name for name, _, _ in SOURCES if not (ROOT / name).exists()]
    if missing:
        raise SystemExit(f"нет исходников: {', '.join(missing)}")
    table = "\n".join(
        f"| `{name}` | **{tag}** | {desc} | `{sha256_of(ROOT / name)}` |"
        for name, tag, desc in SOURCES
    )
    parts = [HEADER.format(sources_table=table)]
    for name, tag, desc in SOURCES:
        parts.append(f"\n# {tag} — {desc}\n\n*Источник: `{name}`*\n")
        parts.append(convert(ROOT / name))
        parts.append("\n---\n")
    return "\n".join(parts).rstrip() + "\n"


def main() -> int:
    check = "--check" in sys.argv
    text = build()
    if check:
        if not OUT.exists():
            print("ТЗ.md отсутствует — запусти scripts/docx_to_tz.py")
            return 1
        current = OUT.read_text(encoding="utf-8")
        # Сверяем только хэши источников: вычитка разметки менять их не должна.
        stale = [sha256_of(ROOT / n) for n, _, _ in SOURCES if sha256_of(ROOT / n) not in current]
        if stale:
            print(f"ТЗ.md устарел: {len(stale)} из {len(SOURCES)} .docx изменились после генерации")
            return 1
        print("ТЗ.md актуален")
        return 0
    OUT.write_text(text, encoding="utf-8")
    print(f"ТЗ.md: {len(text.splitlines())} строк из {len(SOURCES)} .docx")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
