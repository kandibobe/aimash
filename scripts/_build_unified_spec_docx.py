from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "SPEC.md"
OUTPUT = ROOT / "artifacts" / "Aimash_Unified_Technical_Specification_ACCEPTED.docx"

NAVY = "12233F"
BLUE = "1E5AA8"
GREEN = "12805C"
LIGHT_BLUE = "EAF2FB"
LIGHT_GREEN = "E8F5EF"
GRAY = "5F6B7A"
LIGHT_GRAY = "F3F5F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def set_repeat_header(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text.replace("**", "").replace("`", "")


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(clean_inline(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(BLUE)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(clean_inline(text[cursor:]))


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "132033")
    set_cell_margins(cell, 160, 180, 160, 180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(238, 244, 250)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    body = [rows[0], *rows[2:]]
    width = max(len(row) for row in body)
    table = document.add_table(rows=len(body), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_index, row in enumerate(body):
        prevent_row_split(table.rows[r_index])
        for c_index in range(width):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = row[c_index] if c_index < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value.strip())
            for run in p.runs:
                run.font.size = Pt(8)
            if r_index == 0:
                set_cell_shading(cell, NAVY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif r_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 30, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 10.5, GREEN),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "Quote Box" not in styles:
        quote = styles.add_style("Quote Box", WD_STYLE_TYPE.PARAGRAPH)
        quote.base_style = styles["Normal"]
        quote.font.italic = True
        quote.font.color.rgb = RGBColor.from_string(NAVY)
        quote.paragraph_format.left_indent = Cm(0.5)
        quote.paragraph_format.right_indent = Cm(0.3)
        quote.paragraph_format.space_before = Pt(4)
        quote.paragraph_format.space_after = Pt(7)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    r = p.add_run("AIMASH")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    p = document.add_paragraph()
    p.style = document.styles["Title"]
    p.add_run("Единое техническое\nзадание")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("AI-агент управления Google Ads через Telegram")
    r.font.name = "Aptos Display"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    document.add_paragraph("Версия для согласования · 31 июля 2026", style="Subtitle")

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREEN)
    set_cell_margins(cell, 180, 200, 180, 200)
    p = cell.paragraphs[0]
    r = p.add_run("СТАТУС: ЧЕРНОВИК")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p.add_run("\nНе является нормативным источником до явной приёмки заказчиком.")

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.add_run("Базовый принцип\n").bold = True
    r = p.add_run(
        "Hermes автономно думает и готовит действие. Google Ads изменяется только после одного явного подтверждения человека."
    )
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.add_run("Договорная база: ").bold = True
    p.add_run(
        "Aimash_Technical_Specification.docx · Aimash_Flow_Google_Search_4.docx · Информация о клиентах_1.docx"
    )
    document.add_page_break()


def add_contents(document: Document, lines: list[str]) -> None:
    document.add_heading("Содержание", level=1)
    for line in lines:
        if not line.startswith("## "):
            continue
        text = clean_inline(line[3:])
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)
    document.add_page_break()


def build() -> Path:
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    document = Document()
    configure_document(document)
    add_cover(document)
    add_contents(document, lines)

    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([part.strip() for part in lines[i].strip("|").split("|")])
                i += 1
            add_markdown_table(document, rows)
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=3)
        elif line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=2)
        elif line.startswith("> "):
            p = document.add_paragraph(style="Quote Box")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            add_inline(p, line)
        elif line.startswith("- "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            add_inline(p, "• " + line[2:])
        else:
            p = document.add_paragraph()
            add_inline(p, line.strip())
        i += 1

    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "AIMASH  /  ЕДИНОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ  /  ACCEPTED 31.07.2026"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(GRAY)
        add_page_number(section.footer.paragraphs[0])

    core = document.core_properties
    core.title = "Aimash — единое техническое задание"
    core.subject = "AI-агент управления Google Ads через Telegram"
    core.author = "Aimash"
    core.comments = "Черновик для согласования. Нормативным становится только после явной приёмки."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
