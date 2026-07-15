"""Чтение ссылок и файлов в ТЕКСТ для постановки задачи агенту (ingest).

Пользователь даёт ссылку/файл + задачу → бот извлекает текст и передаёт его агенту как
СПРАВОЧНЫЕ ДАННЫЕ (не команды). Любая мутация всё равно проходит confirm-гейт — это backstop
против prompt-injection из стороннего контента (golden rule #8).

Безопасность:
- URL-фетч под SSRF-гардом с ПИННИНГОМ IP: хост резолвится и ВСЕ адреса должны быть публичными
  (приватные/loopback/link-local/reserved/CGNAT/multicast — отказ), после чего коннект идёт РОВНО
  к проверенному IP (`make_ssrf_safe_transport`). Это закрывает TOCTOU DNS-rebinding: проверка и
  соединение — по одному IP, а не по двум резолвам. Редиректы валидируются тем же транспортом
  (резолв на каждый hop). Таймаут + потолок размера ответа.
- Размер файла/текста ограничен (MAX_*). PDF без зависимости pypdf — мягкий отказ.
- Никаких секретов: читаем только то, что прислал пользователь.
"""

from __future__ import annotations

import asyncio
import io
import ipaddress
import re
import socket
from urllib.parse import urlparse

MAX_FETCH_BYTES = 2_000_000  # потолок тела ответа (HTML-страницы лендингов невелики)
MAX_TEXT_CHARS = 8_000  # сколько текста отдаём агенту (ограничивает токены и поверхность инъекции)
MAX_FILE_BYTES = 5_000_000  # потолок принимаемого файла
FETCH_TIMEOUT_S = 10.0
MAX_REDIRECTS = 3
_UA = "AimashBot/1.0 (+google-ads assistant; reads the link you sent)"

_URL_RE = re.compile(r"https?://[^\s<>\"')]+", re.IGNORECASE)
# Текстовые/парсимые расширения файлов (по имени; pdf — отдельным мягким отказом).
_TEXT_EXTS = {"txt", "md", "markdown", "csv", "tsv", "json", "log", "text"}


class IngestError(Exception):
    """Понятная пользователю ошибка чтения (показывается как есть, без сырых трейсбеков)."""


class SSRFBlocked(IngestError):
    """Адрес отклонён SSRF-гардом (внутренний/небезопасный). Подкласс IngestError ⇒ путь
    fetch_url_text показывает его как обычную ошибку чтения; краулер (clients/crawl_fetch) ловит
    его ОТДЕЛЬНО (это МЫ заблокировали, а не сайт упал — предохранитель не трогать)."""


def extract_urls(text: str) -> list[str]:
    """Все http(s)-URL из текста (с обрезкой хвостовой пунктуации .,;)!?»)."""
    out: list[str] = []
    for m in _URL_RE.finditer(text or ""):
        out.append(m.group(0).rstrip(".,;)!?»"))
    return out


def _addr_is_blocked(addr: ipaddress.IPv4Address | ipaddress.IPv6Address) -> bool:
    """Адрес НЕ подлежит запросу (SSRF-риск). `not is_global` — основной гейт: ловит приватные,
    loopback, link-local, а также CGNAT 100.64/10 и shared-address-space, которых нет среди
    боевых сайтов. Явные флаги — для ясной категоризации и на случай version-specific трактовки
    is_global в разных патчах CPython (перестраховка, оба слоя fail-closed)."""
    return (
        not addr.is_global
        or addr.is_private
        or addr.is_loopback
        or addr.is_link_local
        or addr.is_reserved
        or addr.is_multicast
        or addr.is_unspecified
    )


def _validated_connect_ip(ip_str: str) -> str:
    """IP из резолвера → строка IP для коннекта, если он ПУБЛИЧНЫЙ; иначе SSRFBlocked.
    IPv4-mapped IPv6 (`::ffff:a.b.c.d`) разворачиваем в IPv4 и проверяем/коннектимся по нему —
    иначе mapped-loopback (`::ffff:127.0.0.1`) мог бы просочиться там, где флаги на mapped-форме
    трактуются иначе."""
    try:
        addr: ipaddress.IPv4Address | ipaddress.IPv6Address = ipaddress.ip_address(ip_str)
    except ValueError as e:  # мусор/scoped-адрес из резолвера ⇒ fail-closed
        raise SSRFBlocked(f"адрес не распознан: {ip_str}") from e
    mapped = getattr(addr, "ipv4_mapped", None)
    if mapped is not None:
        addr = mapped
    if _addr_is_blocked(addr):
        raise SSRFBlocked(f"адрес заблокирован (внутренний/небезопасный): {ip_str}")
    return str(addr)


def _resolve_pinned_ip(host: str, port: int | None = None) -> str:
    """host → ОДИН публичный IP для коннекта. ВСЕ адреса резолва обязаны быть публичными (политика
    all-public, как в старом _is_public_host) — иначе SSRFBlocked. Из публичных берём
    ДЕТЕРМИНИРОВАННЫЙ (min по (version, packed)): стабильный выбор держит keep-alive пула, т.к.
    транспорт резолвит на каждый запрос. Блокирующий вызов ⇒ звать через asyncio.to_thread.
    Fail-closed: gaierror / пустой ответ / любой приватный адрес ⇒ SSRFBlocked."""
    if not host:
        raise SSRFBlocked("пустой хост")
    try:
        infos = socket.getaddrinfo(host, port, type=socket.SOCK_STREAM)
    except socket.gaierror as e:
        raise SSRFBlocked(f"адрес не разрешается: {host}") from e
    candidates = [_validated_connect_ip(info[4][0]) for info in infos]  # raise на ПЕРВОМ приватном
    if not candidates:
        raise SSRFBlocked(f"адрес не разрешается: {host}")

    def _sort_key(s: str) -> tuple[int, bytes]:
        a = ipaddress.ip_address(s)
        return (a.version, a.packed)

    return min(candidates, key=_sort_key)


def _is_public_host(host: str) -> bool:
    """True, если host резолвится и ВСЕ адреса публичные. Тонкая обёртка над _resolve_pinned_ip
    (единая логика классификации) — оставлена для обратной совместимости и как дешёвый пред-чек.
    АВТОРИТЕТНЫЙ гейт против TOCTOU — сам транспорт (make_ssrf_safe_transport), резолвящий и
    коннектящийся к ОДНОМУ IP. Блокирующий resolve → вызывать через asyncio.to_thread."""
    if not host:
        return False
    try:
        _resolve_pinned_ip(host)
        return True
    except SSRFBlocked:
        return False


_ssrf_transport_cls = (
    None  # ленивый кэш подкласса (httpx импортируется лениво, не на import модуля)
)


def make_ssrf_safe_transport(**transport_kwargs: object):
    """httpx-транспорт, который РЕЗОЛВИТ хост и коннектится к ЗАПИНЕННОМУ публичному IP.

    Закрывает TOCTOU DNS-rebinding: «адрес публичный?» и само соединение идут по ОДНОМУ IP
    (раньше _is_public_host резолвил ОТДЕЛЬНО от httpx-коннекта — между ними DNS мог отдать
    приватный адрес). Резолв — на КАЖДЫЙ запрос, включая редиректы (handle_async_request зовётся
    на каждый hop) ⇒ редирект на внутренний хост тоже отклоняется.

    TLS-идентичность сохранена: для https ставим extension sni_hostname=ИСХОДНЫЙ_ХОСТ — httpcore
    берёт `server_hostname = sni_hostname or origin.host`, поэтому SNI и проверка сертификата идут
    по ИМЕНИ, а TCP-коннект — по IP. Fail-closed: любой отказ резолва ⇒ SSRFBlocked.

    Пиннинг делаем на КОПИИ запроса, входящий объект НЕ мутируем: httpx держит тот же Request и
    резолвит ОТНОСИТЕЛЬНЫЕ редиректы через `request.url.join(Location)` (_client.py). Перепиши мы
    url→IP на исходном объекте — `301 Location: /en/` уехал бы на `https://<IP>/en/` (имя потеряно →
    провал проверки сертификата по голому IP). На копии исходный `request.url` остаётся с ИМЕНЕМ, и
    относительный редирект резолвится на реальный хост (который на след. hop заново пинуется). Копия
    с `stream=` не зовёт `_prepare` ⇒ Host берётся из исходных заголовков (IP в Host не течёт).

    ВНИМАНИЕ: при передаче transport= в httpx.AsyncClient параметры limits=/verify= НА КЛИЕНТЕ
    игнорируются — передавать их СЮДА (в конструктор транспорта)."""
    global _ssrf_transport_cls
    if _ssrf_transport_cls is None:
        import httpx

        class _PinnedIPTransport(httpx.AsyncHTTPTransport):
            async def handle_async_request(self, request: httpx.Request) -> httpx.Response:
                host = request.url.host
                ip = await asyncio.to_thread(_resolve_pinned_ip, host, request.url.port)
                ext = request.extensions
                if request.url.scheme == "https":
                    ext = {**ext, "sni_hostname": host}
                pinned = httpx.Request(
                    request.method,
                    request.url.copy_with(host=ip),  # TCP → IP
                    headers=request.headers,  # Host: реальное имя (порт сохранён)
                    stream=request.stream,  # GET, пустое тело — переиспользование безопасно
                    extensions=ext,
                )
                return await super().handle_async_request(pinned)

        _ssrf_transport_cls = _PinnedIPTransport
    return _ssrf_transport_cls(**transport_kwargs)


# Обвязка страницы (меню/подвал/сайдбар/формы) — она одинакова на всех страницах сайта и раньше
# уезжала в LLM как «текст страницы»: на живом сайте это ровно половина корпуса (8 страниц darial
# отдали по 1431 символа — одно меню). Роли (role=navigation/banner/contentinfo) ловят тему, где
# семантические теги подменены div'ами.
_CHROME_TAGS = ("nav", "footer", "aside", "form")
_CHROME_ROLES = ("navigation", "banner", "contentinfo", "search", "menubar", "complementary")


def _html_to_text(html: str, *, drop_chrome: bool = True) -> str:
    """HTML → читаемый текст: убираем script/style, обвязку страницы, заголовок страницы — в начало.

    drop_chrome=True (дефолт) вырезает nav/footer/aside/form и элементы с навигационными ARIA-ролями.
    <header> сносится только верхнеуровневый: внутри <article>/<main> он несёт H1 самой статьи.
    Отключать имеет смысл лишь там, где меню — и есть искомое содержимое (в проекте таких мест нет)."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    title = (soup.title.string if soup.title and soup.title.string else "").strip()  # ДО decompose
    for tag in soup(["script", "style", "noscript", "template", "svg", "head"]):
        tag.decompose()
    if drop_chrome:
        for tag in soup(list(_CHROME_TAGS)):
            tag.decompose()
        for tag in soup.find_all("header"):
            if tag.find_parent(["article", "main"]) is None:
                tag.decompose()
        for tag in soup.find_all(attrs={"role": True}):
            if str(tag.get("role", "")).strip().lower() in _CHROME_ROLES:
                tag.decompose()
    body = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in body.splitlines()]
    text = "\n".join(ln for ln in lines if ln)
    return (f"{title}\n\n{text}" if title else text).strip()


def _is_textual(content_type: str) -> bool:
    ct = (content_type or "").lower()
    return ct.startswith("text/") or "html" in ct or "xml" in ct or "json" in ct


async def fetch_url_text(url: str) -> str:
    """Прочитать ссылку → текст (HTML → плоский текст). SSRF-гард с пиннингом IP на исходный URL и
    каждый редирект; таймаут; потолок размера. IngestError при отказе/нетекстовом контенте."""
    import httpx

    p = urlparse(url)
    if p.scheme not in ("http", "https"):
        raise IngestError("поддерживаются только http/https-ссылки")

    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            max_redirects=MAX_REDIRECTS,
            timeout=FETCH_TIMEOUT_S,
            headers={"User-Agent": _UA},
            transport=make_ssrf_safe_transport(),
        ) as client:
            async with client.stream("GET", url) as r:
                r.raise_for_status()
                ctype = r.headers.get("content-type", "")
                if not _is_textual(ctype):
                    raise IngestError(f"ссылка не текстовая (тип: {ctype or 'неизвестно'})")
                chunks: list[bytes] = []
                total = 0
                async for chunk in r.aiter_bytes():
                    chunks.append(chunk)
                    total += len(chunk)
                    if total > MAX_FETCH_BYTES:
                        break
                encoding = r.encoding or "utf-8"
            raw = b"".join(chunks)
    except IngestError:
        raise
    except Exception as e:  # сеть/таймаут/HTTP-статус
        raise IngestError(f"не удалось прочитать ссылку ({type(e).__name__})") from e

    try:
        html = raw.decode(encoding, errors="replace")
    except (LookupError, TypeError):
        html = raw.decode("utf-8", errors="replace")
    text = _html_to_text(html) if ("html" in ctype.lower()) else html
    text = text.strip()
    if not text:
        raise IngestError("страница не содержит читаемого текста")
    return text[:MAX_TEXT_CHARS]


def _xlsx_to_text(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out: list[str] = []
    rows = 0
    for ws in wb.worksheets:
        out.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(", ".join(cells))
            rows += 1
            if rows > 2000:  # потолок строк (защита от гигантских книг)
                out.append("…(обрезано)")
                wb.close()
                return "\n".join(out)
    wb.close()
    return "\n".join(out)


def _docx_to_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_file_text(filename: str, data: bytes) -> str:
    """Файл (по расширению) → текст. Поддержаны: txt/md/csv/tsv/json/log, .docx, .xlsx. PDF —
    мягкий отказ (нет зависимости pypdf). IngestError при неподдержанном типе/пустоте/слишком большом."""
    if len(data) > MAX_FILE_BYTES:
        raise IngestError(f"файл слишком большой (> {MAX_FILE_BYTES // 1_000_000} МБ)")
    ext = (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else ""
    if ext == "pdf":
        raise IngestError("PDF пока не поддержан — пришли текст, .docx, .csv или .xlsx")
    try:
        if ext == "docx":
            text = _docx_to_text(data)
        elif ext == "xlsx":
            text = _xlsx_to_text(data)
        elif ext in _TEXT_EXTS or ext == "":
            text = data.decode("utf-8", errors="replace")
        else:
            raise IngestError(f"тип файла .{ext} не поддержан — пришли txt/csv/json/.docx/.xlsx")
    except IngestError:
        raise
    except Exception as e:  # битый файл/неверный формат
        raise IngestError(f"не удалось прочитать файл ({type(e).__name__})") from e
    text = (text or "").strip()
    if not text:
        raise IngestError("файл пустой или без читаемого текста")
    return text[:MAX_TEXT_CHARS]
