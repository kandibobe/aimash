"""Выгрузка /audit в Word (.docx) — reports/docx.py.

Инварианты (каждый — про честность/деньги, не про красоту):
• порядок секций тот же, что у .xlsx/Sheets: Находки ПЕРВОЙ → По семьям → Обзор ПОСЛЕДНИМ
  (файл открывают ради находок, «Обзор» — копия поста в чате; замечание 5, 2026-07-17);
• проза находок = проза карточки (тот же реестр reports.findings — расходиться нечему),
  worst-first;
• экспорт — БУМАГА: ни колонки suggested_operation, ни кнопки «применить» (золотое правило №3);
• деньги в ячейках — СТРОКОЙ с двумя знаками (у docx нет числового формата), пусто остаётся пусто;
• safe_row НЕ применяется ОСОЗНАННО: Word не исполняет текст ячейки как формулу, апостроф-префикс
  Excel был бы буквальным мусором — имя `=HYPERLINK(...)` лежит как есть, БЕЗ ведущего апострофа;
• EN-документ без кириллицы (RU-утечка в артефакте);
• write_audit_docx пишет валидный .docx (открывается обратно).
"""

from __future__ import annotations

import pathlib
import sys
from datetime import date
from types import SimpleNamespace

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1]))

import pytest  # noqa: E402
from docx import Document  # noqa: E402

from audit.engine import build_audit  # noqa: E402
from reports.docx import build_audit_docx, write_audit_docx  # noqa: E402
from reports.findings import findings_rows  # noqa: E402
from reports.period import last_n_days  # noqa: E402
from reports.queries import Breakdown, Metrics  # noqa: E402

ROOT = pathlib.Path(__file__).resolve().parents[1]
EVIL = '=HYPERLINK("http://evil","click")'
_CYR = set("абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ")


def _report(campaign: str = "Search Brand"):
    """Аккаунт с расходом и БЕЗ конверсий — движку есть что сказать (иначе тест зелен вхолостую)."""
    m = Metrics(impressions=1000, clicks=100, cost_micros=500_000_000, conversions=0.0)
    camp = Breakdown(
        key="campaign",
        title="Кампании",
        dim_headers=["Кампания", "Статус"],
        rows=[((campaign, "ENABLED"), m)],
    )
    return SimpleNamespace(
        customer_id="7753643025",
        totals=m,
        prev_totals=m,
        period=last_n_days(7, today=date(2026, 6, 25)),
        currency="USD",
        breakdowns=[camp],
    )


def _result(campaign: str = "Search Brand"):
    r = build_audit(_report(campaign))
    assert r.findings and r.families, "движку есть что сказать про этот аккаунт"
    return r


def _cells(doc) -> list[str]:
    return [c.text for t in doc.tables for row in t.rows for c in row.cells]


def _headings(doc, style: str = "Heading 1") -> list[str]:
    return [p.text for p in doc.paragraphs if p.style.name == style]


def test_docx_sections_mirror_the_xlsx_order():
    """Тот же порядок, что закреплён для .xlsx/Sheets: Находки → По семьям → Обзор (последним)."""
    doc = build_audit_docx(_result(), "ru")
    assert _headings(doc) == ["Находки", "По семьям", "Обзор"]
    assert _headings(doc, "Title") == ["Аудит аккаунта 7753643025"]
    assert len(doc.tables) == 2  # Находки + По семьям; Обзор — проза, НЕ таблица


def test_docx_speaks_the_same_prose_as_the_audit_card():
    """Документ — не пересказ: проза каждой находки и её check_id из ЕДИНОГО реестра, worst-first;
    обзор = карточка /audit (all_quick_wins-бумага) параграфами."""
    from audit.render import finding_text, render_audit

    result = _result()
    doc = build_audit_docx(result, "ru")
    cells = _cells(doc)
    for f in result.findings:
        assert finding_text(f, "ru", result.currency) in cells
    findings_tbl = doc.tables[0]
    last_col = [row.cells[-1].text for row in findings_tbl.rows[1:]]
    assert last_col == [f.check_id for f in result.findings]  # порядок карточки
    paras = [p.text for p in doc.paragraphs]
    for line in render_audit(result, "ru", actions=False, all_quick_wins=True).split("\n"):
        if line:
            assert line in paras, f"обзор карточки не доехал до документа: {line}"


def test_docx_is_paper_not_a_button():
    """Золотое правило №3: экспорт не предлагает применить. Ни поля операции, ни кнопки — по коду."""
    src = (ROOT / "reports" / "docx.py").read_text(encoding="utf-8")
    assert "suggested_operation" not in src
    assert "ONE_TAP_OPS" not in src and "one_tap" not in src
    result = _result()
    ops = {getattr(f, "suggested_operation", None) for f in result.findings}
    assert ops - {None}, "фикстура без быстрых побед не проверяет запрет (нужна хоть одна)"
    assert not (set(_cells(build_audit_docx(result, "ru"))) & (ops - {None}))


def test_docx_en_has_no_cyrillic():
    doc = build_audit_docx(_result(), "en")
    texts = [p.text for p in doc.paragraphs] + _cells(doc)
    leaks = [t for t in texts if _CYR & set(t)]
    assert not leaks, f"RU-утечка в EN-документе: {leaks}"
    assert _headings(doc) == ["Findings", "By family", "Overview"]


def test_docx_money_is_text_and_names_are_literal():
    """Деньги — строкой с 2 знаками (нет данных → пусто, не «0.00»); имя `=HYPERLINK…` лежит
    БУКВАЛЬНО, без апострофа safe_row: Word не исполняет ячейку как формулу, а «'=…» в тексте —
    мусор."""
    result = _result(campaign=EVIL)
    doc = build_audit_docx(result, "ru")
    findings_tbl = doc.tables[0]
    body = [row.cells for row in findings_tbl.rows[1:]]
    for cells_row, src_row in zip(body, findings_rows(result, "ru"), strict=True):
        for cell, v in zip(cells_row, src_row, strict=True):
            assert isinstance(cell.text, str)
            if isinstance(v, float):
                assert cell.text == f"{v:.2f}", "деньги в docx — строка с двумя знаками"
            elif v == "":
                assert cell.text == ""
    cells = _cells(doc)
    assert EVIL in cells, "имя кампании должно лежать буквально (ячейка == имени)"
    assert not any(t.startswith("'=") for t in cells), "апостроф safe_row в Word — буквальный мусор"


def test_write_audit_docx_roundtrip(tmp_path):
    """Файл на диске — валидный .docx: открывается обратно, секции и таблицы на месте."""
    path = str(tmp_path / "audit.docx")
    assert write_audit_docx(_result(), path, "ru") == path
    doc = Document(path)
    assert _headings(doc) == ["Находки", "По семьям", "Обзор"]
    assert len(doc.tables) == 2


if __name__ == "__main__":  # pragma: no cover
    pytest.main([__file__, "-q"])
