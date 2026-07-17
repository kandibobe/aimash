"""Экспорт результата /audit в Word (.docx). READ-ONLY, без секретов, GR3: бумага.

Word — документ для ЧТЕНИЯ (отдать клиенту/руководителю), а не таблица для работы: три секции —
«Находки» (таблица всех находок worst-first), «По семьям» (свод) и «Обзор» (карточка прозой,
ПОСЛЕДНИМ — это копия поста в чате, замечание 5, 2026-07-17). Секцию сырых ДАННЫХ (Сводка/разбивки/
запросы) сюда осознанно НЕ кладём: в документе её нельзя ни сортировать, ни фильтровать — для
работы с данными есть .xlsx и Google Sheets (build_audit_workbook / build_audit_sheets_data).

Раскладку строк даёт reports.findings — ТА ЖЕ, что у .xlsx и Sheets (один реестр, расходиться
нечему); здесь только запись в docx-API. Деньги пишутся СТРОКОЙ с двумя знаками (у ячейки docx
нет числового формата), пустой факт — пустой ячейкой (нет данных ≠ ноль).

⚠️ safe_row здесь НЕ применяем — осознанно: Word не исполняет текст ячейки как формулу (инъекция
`=HYPERLINK(...)` — угроза Excel/Sheets), а апостроф-префикс Excel в Word был бы буквальным
мусором в тексте. Ни колонки операции, ни кнопки «применить» нет и не будет: экспорт — бумага
(золотое правило №3), гард — tests/test_audit_docx.py::test_docx_is_paper_not_a_button.
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_ORIENT

from reports.findings import (
    FAMILY_SUMMARY_TITLE,
    FINDINGS_TITLE,
    OVERVIEW_TITLE,
    family_summary_headers,
    family_summary_rows,
    findings_headers,
    findings_meta_rows,
    findings_rows,
)
from reports.labels import loc


def _cell_text(v) -> str:
    """Ячейка docx — всегда строка: float (деньги/штраф, уже round(...,2) в реестре) — с двумя
    знаками, остальное как есть. Пустое остаётся пустым (нет данных ≠ ноль)."""
    if v is None or v == "":
        return ""
    if isinstance(v, float):
        return f"{v:.2f}"
    return str(v)


def _add_table(doc, headers: list[str], rows: list[list]) -> None:
    """Таблица «шапка + строки» стилем Table Grid (есть в дефолтном шаблоне), шапка жирным."""
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = _cell_text(v)


def build_audit_docx(result, lang: str = "ru"):
    """Word-документ выгрузки /audit из УЖЕ посчитанного AuditResult — без сети и без доп-чтений
    Google Ads (аудит не пересобираем, result из кэша bot-слоя). lang — язык ПОДПИСЕЙ
    (reports.labels); значения ячеек (имена кампаний, тексты находок) не трогаем."""
    doc = Document()
    # Портрет не вмещает колонки находок — альбомная. Смена orientation размеры страницы
    # НЕ свапает (python-docx делает это только флагом) — меняем ширину/высоту явно.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    cid = getattr(result, "customer_id", "") or ""
    doc.add_heading(f"Account audit {cid}" if lang == "en" else f"Аудит аккаунта {cid}", 0)
    currency = getattr(result, "currency", "") or ""

    doc.add_heading(loc(FINDINGS_TITLE, lang), level=1)
    _add_table(doc, findings_headers(currency, lang), findings_rows(result, lang))

    doc.add_heading(loc(FAMILY_SUMMARY_TITLE, lang), level=1)
    _add_table(doc, family_summary_headers(currency, lang), family_summary_rows(result, lang))

    doc.add_heading(loc(OVERVIEW_TITLE, lang), level=1)
    for row in findings_meta_rows(result, lang):  # проза карточки, по строке на параграф
        doc.add_paragraph(row[0] if row else "")
    return doc


def write_audit_docx(result, path: str, lang: str = "ru") -> str:
    """Сохранить выгрузку /audit в .docx по пути path. Возвращает path. lang — язык ПОДПИСЕЙ."""
    build_audit_docx(result, lang).save(path)
    return path
